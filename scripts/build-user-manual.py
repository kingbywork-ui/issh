from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "user-manual" / "iSSH-使用手册.md"
OUTPUT = ROOT / "docs" / "user-manual" / "iSSH-使用手册.docx"

FONT = "Microsoft YaHei"
FONT_EAST_ASIA = "微软雅黑"
NAVY = "173B57"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6570"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_w = table_pr.find(qn("w:tblW"))
    if table_w is None:
        table_w = OxmlElement("w:tblW")
        table_pr.append(table_w)
    table_w.set(qn("w:w"), str(sum(widths)))
    table_w.set(qn("w:type"), "dxa")

    table_ind = table_pr.find(qn("w:tblInd"))
    if table_ind is None:
        table_ind = OxmlElement("w:tblInd")
        table_pr.append(table_ind)
    table_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    tail = paragraph.add_run(" 页")
    set_font(tail, size=9, color=MUTED)


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    style_num_pr = doc.styles["List Number"].element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    abstract_num_id = None
    existing_ids = []
    for num in numbering.findall(qn("w:num")):
        num_id = int(num.get(qn("w:numId")))
        existing_ids.append(num_id)
        if num_id == base_num_id:
            abstract = num.find(qn("w:abstractNumId"))
            abstract_num_id = int(abstract.get(qn("w:val")))
    if abstract_num_id is None:
        raise RuntimeError("List Number style has no numbering definition")

    new_num_id = max(existing_ids, default=0) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    new_num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)
    return new_num_id


def add_numbered_paragraph(doc, text, num_id):
    paragraph = doc.add_paragraph(style="List Number")
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
    num_pr.get_or_add_numId().set(qn("w:val"), str(num_id))
    add_inline(paragraph, text)
    return paragraph


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (11.5, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)


def add_inline(paragraph, text):
    token_pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    lead = p.add_run("说明  ")
    set_font(lead, bold=True, color=DARK_BLUE)
    add_inline(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def table_widths(headers):
    if len(headers) == 3:
        return [1600, 4500, 3260]
    first = headers[0] if headers else ""
    if "操作" in first:
        return [5000, 4360]
    if "Codex Desktop" in first:
        return [3100, 6260]
    return [2500, 6860]


def add_markdown_table(doc, rows):
    headers = rows[0]
    body = rows[2:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = table_widths(headers)
    set_table_geometry(table, widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    header_properties.append(table_header)

    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        set_font(run, size=9.5, bold=True, color=NAVY)

    for row_values in body:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            if len(body) > 4 and len(rows) % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.15
            add_inline(p, value)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(9.3)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_image(doc, alt, relative_path):
    image_path = SOURCE.parent / relative_path
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    if image_path.name == "02-settings.png":
        width = Inches(5.35)
    elif "codex" in image_path.name.lower():
        width = Inches(4.6)
    else:
        width = Inches(6.2)
    inline_shape = run.add_picture(str(image_path), width=width)
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt)
    caption = doc.add_paragraph(f"图：{alt}", style="Caption")
    caption.paragraph_format.keep_with_next = False


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(row)
        index += 1
    return rows, index


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("用户操作与安全接入指南")
    set_font(run, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("iSSH 使用手册")
    set_font(run, size=30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    run = subtitle.add_run("AI 增强型 SSH 终端")
    set_font(run, size=15, color=DARK_BLUE)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_after = Pt(85)
    run = version.add_run("适用版本 0.0.7 · Windows x64")
    set_font(run, size=10.5, color=MUTED)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.65)
    summary.paragraph_format.right_indent = Inches(0.65)
    summary.paragraph_format.space_after = Pt(30)
    run = summary.add_run("从首次连接服务器，到 AI 命令补全和 Codex Desktop MCP 接入，一份面向普通用户的完整指南。")
    set_font(run, size=12, color=NAVY)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("2026 年 7 月")
    set_font(run, size=10.5, color=MUTED)
    date.add_run().add_break(WD_BREAK.PAGE)


def add_navigation(doc, headings):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("内容导航")
    set_font(run, size=20, bold=True, color=NAVY)

    intro = doc.add_paragraph("本手册按实际使用顺序组织，可从对应章节直接开始阅读。")
    intro.paragraph_format.space_after = Pt(14)

    navigation_num_id = create_numbering_instance(doc)
    for heading in headings:
        text = re.sub(r"^\d+\.\s*", "", heading)
        add_numbered_paragraph(doc, text, navigation_num_id)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_styles(doc)
    doc.core_properties.title = "iSSH 使用手册"
    doc.core_properties.subject = "iSSH 0.0.7 用户操作与安全接入指南"
    doc.core_properties.author = "iSSH"
    doc.core_properties.keywords = "iSSH, SSH, 终端, 命令补全, CLI, MCP, Codex"

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run("iSSH 使用手册  |  版本 0.0.7")
    set_font(run, size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    add_navigation(doc, headings)

    index = 0
    active_numbering_id = None
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        is_numbered_line = bool(re.match(r"^\d+\.\s+", line))
        if not is_numbered_line:
            active_numbering_id = None

        if index == 0 and line.startswith("# "):
            index += 1
            continue
        if line.startswith("**适用版本") or line.startswith("**手册日期"):
            index += 1
            continue
        if line == "---" or not line:
            index += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                add_image(doc, match.group(1), match.group(2))
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            if len(rows) >= 2:
                add_markdown_table(doc, rows)
            continue
        if line.startswith(">"):
            add_callout(doc, line[1:].strip())
            index += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if is_numbered_line:
            text = re.sub(r"^\d+\.\s+", "", line)
            if active_numbering_id is None:
                active_numbering_id = create_numbering_instance(doc)
            add_numbered_paragraph(doc, text, active_numbering_id)
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            index += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line)
        next_index = index + 1
        while next_index < len(lines) and not lines[next_index].strip():
            next_index += 1
        if next_index < len(lines) and lines[next_index].strip().startswith("|"):
            p.paragraph_format.keep_with_next = True
        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
